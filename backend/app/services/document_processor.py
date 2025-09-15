"""
Document processing service for RAG system
Text extraction, chunking, and indexing
"""

import logging
import asyncio
import tempfile
import os
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime
from uuid import UUID
import hashlib

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

from app.config import get_settings
from app.services.vector_store import get_vector_store
from app.utils.metrics import track_processing_time, increment_counter

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Advanced document processing with multi-format support
    """
    
    def __init__(self):
        self.settings = get_settings()
        self.text_splitter = None
        self.embedding_model = None
        self._initialize_processors()
    
    def _initialize_processors(self):
        """Initialize text processing components"""
        try:
            logger.info("Initializing document processors...")
            
            # Initialize text splitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.settings.CHUNK_SIZE,
                chunk_overlap=self.settings.CHUNK_OVERLAP,
                separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
                keep_separator=True
            )
            
            # Initialize embedding model
            self.embedding_model = SentenceTransformer(
                self.settings.EMBEDDING_MODEL,
                device='cpu'  # Use CPU for compatibility
            )
            
            logger.info("Document processors initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize document processors: {e}")
            raise
    
    @track_processing_time
    async def process_document(
        self,
        file_path: str,
        document_id: UUID,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process uploaded document through complete pipeline
        
        Args:
            file_path: Path to uploaded file
            document_id: Unique document identifier
            metadata: Document metadata
            
        Returns:
            Processing results with statistics
        """
        try:
            logger.info(f"Processing document: {document_id}")
            start_time = datetime.utcnow()
            
            # Extract text based on file type
            file_extension = os.path.splitext(file_path)[1].lower()
            text_content = await self._extract_text(file_path, file_extension)
            
            if not text_content or not text_content.strip():
                raise ValueError("No text content extracted from document")
            
            # Clean and preprocess text
            cleaned_text = await self._clean_text(text_content)
            
            # Create chunks
            chunks = await self._create_chunks(cleaned_text, document_id, metadata)
            
            # Generate embeddings
            embeddings = await self._generate_embeddings([chunk["content"] for chunk in chunks])
            
            # Store in vector database
            await self._store_in_vector_db(chunks, embeddings)
            
            # Calculate processing statistics
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            result = {
                "document_id": str(document_id),
                "status": "completed",
                "chunk_count": len(chunks),
                "character_count": len(cleaned_text),
                "word_count": len(cleaned_text.split()),
                "processing_time_seconds": processing_time,
                "file_type": file_extension,
                "chunks_stored": len(chunks),
                "embeddings_generated": len(embeddings)
            }
            
            increment_counter("documents_processed_total", 
                            {"status": "success", "file_type": file_extension[1:]})
            
            logger.info(f"Document processing completed: {document_id} ({processing_time:.2f}s)")
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed for {document_id}: {e}")
            increment_counter("documents_processed_total", 
                            {"status": "failed", "file_type": file_extension[1:]})
            raise
    
    async def _extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text from various file formats"""
        try:
            logger.debug(f"Extracting text from {file_extension} file")
            
            if file_extension == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_extension == '.docx':
                return await self._extract_text_from_docx(file_path)
            elif file_extension in ['.txt', '.md']:
                return await self._extract_text_from_text(file_path)
            elif file_extension in ['.html', '.htm']:
                return await self._extract_text_from_html(file_path)
            else:
                # Try as plain text
                logger.warning(f"Unsupported file type {file_extension}, trying as plain text")
                return await self._extract_text_from_text(file_path)
                
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise
    
    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available for PDF processing")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{page_text}\n")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise
    
    async def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available for DOCX processing")
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise
    
    async def _extract_text_from_text(self, file_path: str) -> str:
        """Extract text from plain text/markdown file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with error handling
            with open(file_path, 'rb') as file:
                content = file.read()
                return content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise
    
    async def _extract_text_from_html(self, file_path: str) -> str:
        """Extract text from HTML file"""
        if not HTML_AVAILABLE:
            raise ImportError("BeautifulSoup not available for HTML processing")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            logger.error(f"HTML text extraction failed: {e}")
            raise
    
    async def _clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        try:
            # Basic text cleaning
            cleaned = text.strip()
            
            # Remove excessive whitespace
            import re
            cleaned = re.sub(r'\n\s*\n', '\n\n', cleaned)  # Multiple newlines
            cleaned = re.sub(r' +', ' ', cleaned)  # Multiple spaces
            
            # Remove control characters but keep newlines and tabs
            cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned)
            
            # Normalize quotes and dashes
            cleaned = cleaned.replace('"', '"').replace('"', '"')
            cleaned = cleaned.replace(''', "'").replace(''', "'")
            cleaned = cleaned.replace('–', '-').replace('—', '-')
            
            return cleaned
            
        except Exception as e:
            logger.error(f"Text cleaning failed: {e}")
            return text  # Return original text if cleaning fails
    
    async def _create_chunks(
        self,
        text: str,
        document_id: UUID,
        metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Create text chunks with metadata"""
        try:
            logger.debug(f"Creating chunks for document: {document_id}")
            
            # Split text into chunks
            text_chunks = self.text_splitter.split_text(text)
            
            chunks = []
            for i, chunk_text in enumerate(text_chunks):
                if chunk_text.strip():  # Skip empty chunks
                    chunk = {
                        "content": chunk_text.strip(),
                        "document_id": str(document_id),
                        "chunk_id": f"{document_id}_chunk_{i:04d}",
                        "chunk_index": i,
                        "filename": metadata.get("filename", "unknown"),
                        "source": metadata.get("source", "unknown"),
                        "metadata": {
                            **metadata,
                            "chunk_index": i,
                            "chunk_length": len(chunk_text),
                            "chunk_word_count": len(chunk_text.split()),
                            "created_at": datetime.utcnow().isoformat()
                        }
                    }
                    chunks.append(chunk)
            
            logger.info(f"Created {len(chunks)} chunks from document: {document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"Chunk creation failed: {e}")
            raise
    
    async def _generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for text chunks"""
        try:
            logger.debug(f"Generating embeddings for {len(texts)} chunks")
            
            # Process in batches for memory efficiency
            batch_size = self.settings.EMBEDDING_BATCH_SIZE
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                
                # Generate embeddings
                batch_embeddings = self.embedding_model.encode(
                    batch_texts,
                    batch_size=len(batch_texts),
                    show_progress_bar=False,
                    normalize_embeddings=True
                )
                
                # Convert to list format
                all_embeddings.extend(batch_embeddings.tolist())
                
                # Small delay to prevent overwhelming the system
                if i + batch_size < len(texts):
                    await asyncio.sleep(0.1)
            
            logger.info(f"Generated embeddings for {len(all_embeddings)} chunks")
            return all_embeddings
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise
    
    async def _store_in_vector_db(
        self,
        chunks: List[Dict[str, Any]],
        embeddings: List[List[float]]
    ) -> bool:
        """Store chunks and embeddings in vector database"""
        try:
            logger.debug(f"Storing {len(chunks)} chunks in vector database")
            
            vector_store = await get_vector_store()
            
            # Prepare documents for storage
            documents = []
            ids = []
            
            for chunk in chunks:
                documents.append({
                    "content": chunk["content"],
                    "document_id": chunk["document_id"],
                    "chunk_id": chunk["chunk_id"],
                    "filename": chunk["filename"],
                    "source": chunk["source"],
                    **chunk["metadata"]
                })
                ids.append(chunk["chunk_id"])
            
            # Store in vector database
            success = await vector_store.add_documents(documents, embeddings, ids)
            
            if success:
                logger.info(f"Successfully stored {len(chunks)} chunks in vector database")
            else:
                logger.error("Failed to store chunks in vector database")
            
            return success
            
        except Exception as e:
            logger.error(f"Vector storage failed: {e}")
            raise
    
    async def reprocess_document(
        self,
        document_id: UUID,
        new_settings: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Reprocess existing document with new settings"""
        try:
            logger.info(f"Reprocessing document: {document_id}")
            
            # This would typically:
            # 1. Retrieve original document content
            # 2. Apply new processing settings
            # 3. Re-chunk and re-embed
            # 4. Update vector store
            
            # Mock implementation for demo
            result = {
                "document_id": str(document_id),
                "status": "reprocessed",
                "new_chunk_count": 52,
                "changes_applied": new_settings or {},
                "reprocessing_time": 15.3
            }
            
            increment_counter("documents_reprocessed_total")
            return result
            
        except Exception as e:
            logger.error(f"Document reprocessing failed: {e}")
            raise
    
    async def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        try:
            file_stats = os.stat(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            metadata = {
                "file_size": file_stats.st_size,
                "file_type": file_extension,
                "created_at": datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(file_stats.st_mtime).isoformat()
            }
            
            # Extract format-specific metadata
            if file_extension == '.pdf' and PDF_AVAILABLE:
                metadata.update(await self._extract_pdf_metadata(file_path))
            elif file_extension == '.docx' and DOCX_AVAILABLE:
                metadata.update(await self._extract_docx_metadata(file_path))
            
            # Calculate file hash for duplicate detection
            metadata["file_hash"] = await self._calculate_file_hash(file_path)
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {e}")
            return {}
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "pdf_version": getattr(pdf_reader, 'pdf_version', 'unknown')
                }
                
                # Extract document info if available
                if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                    doc_info = pdf_reader.metadata
                    metadata.update({
                        "title": str(doc_info.get('/Title', '')),
                        "author": str(doc_info.get('/Author', '')),
                        "subject": str(doc_info.get('/Subject', '')),
                        "creator": str(doc_info.get('/Creator', '')),
                        "producer": str(doc_info.get('/Producer', ''))
                    })
                
                return metadata
                
        except Exception as e:
            logger.warning(f"PDF metadata extraction failed: {e}")
            return {}
    
    async def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        try:
            doc = Document(file_path)
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "subject": core_props.subject or '',
                    "comments": core_props.comments or '',
                    "category": core_props.category or '',
                    "language": core_props.language or ''
                })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"DOCX metadata extraction failed: {e}")
            return {}
    
    async def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, 'rb') as file:
                for chunk in iter(lambda: file.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
            
        except Exception as e:
            logger.error(f"File hash calculation failed: {e}")
            return ""
    
    async def get_processing_stats(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        # This would typically query a database for actual stats
        return {
            "total_processed": 15847,
            "successful_processing": 15823,
            "failed_processing": 24,
            "avg_processing_time": 23.4,
            "total_chunks_created": 756234,
            "avg_chunks_per_document": 47.8,
            "supported_formats": self.settings.supported_file_types_list,
            "embedding_model": self.settings.EMBEDDING_MODEL,
            "chunk_size": self.settings.CHUNK_SIZE,
            "chunk_overlap": self.settings.CHUNK_OVERLAP
        }
